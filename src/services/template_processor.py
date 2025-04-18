import re
from docx import Document
from typing import List, Dict
from src.models.template import Template

class TemplateProcessor:
    """Service for processing document templates"""
    
    def __init__(self):
        self.variable_pattern = r'\{\{([^}]+)\}\}'  # Matches {{variable_name}}
    
    def extract_variables(self, template: Template) -> List[str]:
        """Extract variables from a template document"""
        doc = Document(template.path)
        variables = set()
        
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            matches = re.finditer(self.variable_pattern, paragraph.text)
            for match in matches:
                variables.add(match.group(1).strip())
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.finditer(self.variable_pattern, cell.text)
                    for match in matches:
                        variables.add(match.group(1).strip())
        
        return list(variables)
    
    def process_template(self, template: Template, data: Dict[str, str]) -> Document:
        """Process a template with provided data"""
        doc = Document(template.path)
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            self._replace_variables(paragraph, data)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_variables(paragraph, data)
        
        return doc
    
    def _replace_variables(self, paragraph, data: Dict[str, str]) -> None:
        """Replace variables in a paragraph with provided data"""
        text = paragraph.text
        for var_name, value in data.items():
            pattern = r'\{\{' + re.escape(var_name) + r'\}\}'
            text = re.sub(pattern, value, text)
        
        if text != paragraph.text:
            paragraph.text = text 