import uuid
from fastapi import UploadFile
from typing import List, Dict, Any
import PyPDF2
import docx
import os
import tempfile
from openai import AsyncOpenAI
import tiktoken
import numpy as np
from models.schemas import ProcessedChunk, ProcessedDocument, FileStatus
from services.vector_store import VectorStore

# Initialize OpenAI client
openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

async def process_document(file: UploadFile, user_id: str) -> str:
    """
    Process a document file and return the file ID.
    """
    file_id = str(uuid.uuid4())
    
    try:
        # Extract text from the document
        text = await extract_text(file)
        
        # Create chunks from the text
        chunks = await create_chunks(text)
        
        # Generate embeddings for chunks
        chunk_objects = await generate_chunk_embeddings(chunks, file_id, user_id)
        
        # Create processed document object
        document = ProcessedDocument(
            file_id=file_id,
            filename=file.filename,
            chunks=chunk_objects,
            metadata={
                "file_type": os.path.splitext(file.filename)[1].lower(),
                "original_size": len(text),
                "chunk_count": len(chunks)
            }
        )
        
        # Store document and chunks in the database
        await store_document(document, user_id)
        
        return file_id
        
    except Exception as e:
        # Update document status with error
        await VectorStore.update_document_status(file_id, "error", str(e))
        raise e

async def extract_text(file: UploadFile) -> str:
    """
    Extract text from a PDF or DOCX file
    """
    content = ""
    file_extension = os.path.splitext(file.filename)[1].lower()
    
    # Create a temporary file to save the uploaded content
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
        temp_file.write(await file.read())
        temp_file_path = temp_file.name
    
    try:
        if file_extension == '.pdf':
            with open(temp_file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
        
        elif file_extension == '.docx':
            doc = docx.Document(temp_file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"
        
        else:
            content = "Unsupported file format"
    
    finally:
        # Cleanup the temporary file
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
        
        # Reset file pointer for potential reuse
        await file.seek(0)
    
    return content

async def create_chunks(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
    """
    Split text into chunks with overlap using token-based splitting
    """
    # Initialize tokenizer
    tokenizer = tiktoken.get_encoding("cl100k_base")
    
    # Get tokens
    tokens = tokenizer.encode(text)
    
    # Calculate chunk sizes in tokens
    chunk_size_tokens = chunk_size
    overlap_tokens = overlap
    
    # Create chunks
    chunks = []
    for i in range(0, len(tokens), chunk_size_tokens - overlap_tokens):
        chunk_tokens = tokens[i:i + chunk_size_tokens]
        chunk_text = tokenizer.decode(chunk_tokens)
        chunks.append(chunk_text.strip())
    
    return chunks

async def generate_chunk_embeddings(
    chunks: List[str],
    document_id: str,
    user_id: str
) -> List[ProcessedChunk]:
    """
    Generate embeddings for a list of text chunks using OpenAI's API
    """
    processed_chunks = []
    
    for i, chunk in enumerate(chunks):
        try:
            # Generate embedding
            response = await openai_client.embeddings.create(
                input=chunk,
                model="text-embedding-3-small"
            )
            embedding = response.data[0].embedding
            
            # Create chunk object
            chunk_obj = ProcessedChunk(
                chunk_id=str(uuid.uuid4()),
                content=chunk,
                embedding=embedding,
                metadata={
                    "document_id": document_id,
                    "chunk_number": i,
                    "token_count": len(tiktoken.get_encoding("cl100k_base").encode(chunk))
                }
            )
            
            processed_chunks.append(chunk_obj)
            
        except Exception as e:
            print(f"Error generating embedding for chunk {i}: {str(e)}")
            # Continue processing other chunks
            continue
    
    return processed_chunks

async def store_document(document: ProcessedDocument, user_id: str):
    """
    Store the processed document and its chunks in the database
    """
    # Store document metadata
    await VectorStore.store_document(
        file_id=document.file_id,
        user_id=user_id,
        filename=document.filename,
        metadata=document.metadata
    )
    
    try:
        # Store chunks in vector database
        for chunk in document.chunks:
            await VectorStore.store_chunk(chunk, user_id, document.file_id)
        
        # Update document status to completed
        await VectorStore.update_document_status(document.file_id, "completed")
        
    except Exception as e:
        await VectorStore.update_document_status(document.file_id, "error", str(e))
        raise e

async def get_processing_status(file_id: str, user_id: str) -> FileStatus:
    """
    Get the processing status of a file from the database
    """
    return await VectorStore.get_document_status(file_id, user_id) 