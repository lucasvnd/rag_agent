from typing import Dict, Optional, Any
from docx import Document
from pathlib import Path
import re
import json
from models.templates import Template
from fastapi import HTTPException
from docxtpl import DocxTemplate
import aiofiles
import os

class TemplateProcessor:
    """Service for processing DOCX templates and handling variable replacement"""
    
    def __init__(self):
        self.variable_pattern = re.compile(r'\{\{([^}]+)\}\}')  # Matches {{variable_name}}
    
    async def load_template(self, template: Template) -> Document:
        """Load a DOCX template file"""
        try:
            template_path = Path(template.file_path)
            if not template_path.exists():
                raise HTTPException(status_code=404, detail=f"Template file not found: {template.name}")
            return Document(template_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error loading template: {str(e)}")
    
    async def get_template_variables(self, template: Template) -> Dict[str, str]:
        """Extract all variables from a template"""
        doc = await self.load_template(template)
        variables = set()
        
        # Search for variables in paragraphs
        for paragraph in doc.paragraphs:
            matches = self.variable_pattern.finditer(paragraph.text)
            variables.update(match.group(1).strip() for match in matches)
        
        # Search for variables in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = self.variable_pattern.finditer(paragraph.text)
                        variables.update(match.group(1).strip() for match in matches)
        
        return {var: "" for var in variables}
    
    async def validate_template_data(self, template: Template, data: Dict[str, Any]) -> bool:
        """Validate that all required template variables are provided"""
        required_vars = await self.get_template_variables(template)
        missing_vars = [var for var in required_vars if var not in data]
        
        if missing_vars:
            raise HTTPException(
                status_code=400,
                detail=f"Missing required variables: {', '.join(missing_vars)}"
            )
        return True
    
    async def process_template(self, template: Template, data: Dict[str, Any]) -> DocxTemplate:
        """
        Process a template with the provided data
        
        Args:
            template: Template object containing template information
            data: Dictionary of data to fill the template with
            
        Returns:
            DocxTemplate: Processed document
        """
        if not os.path.exists(template.file_path):
            raise FileNotFoundError(f"Template file not found: {template.file_path}")
            
        # Load the template
        doc = DocxTemplate(template.file_path)
        
        # Render the template with the provided data
        doc.render(data)
        
        return doc
    
    async def save_processed_template(self, doc: DocxTemplate, output_path: str) -> None:
        """
        Save a processed template to disk
        
        Args:
            doc: Processed DocxTemplate object
            output_path: Path where to save the document
        """
        # Ensure the output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Save the document
        doc.save(output_path)
    
    async def update_template_metadata(self, template: Template, variables: Dict[str, str]) -> Dict:
        """Update template metadata with extracted variables"""
        metadata = template.metadata or {}
        metadata["variables"] = variables
        return metadata 