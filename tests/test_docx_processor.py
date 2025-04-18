import pytest
from pathlib import Path
import uuid
from datetime import datetime
from typing import Dict, Any, List
from docx import Document as DocxDocument

from services.document_processor import DocumentProcessor
from models.documents import Document, DocumentStatus

@pytest.fixture
def sample_docx_content():
    """Create a sample DOCX document for testing."""
    doc = DocxDocument()
    doc.add_paragraph("Test Content")
    doc.add_paragraph("Second paragraph")
    return doc

@pytest.fixture
def sample_docx_file(temp_dir: Path, sample_docx_content):
    """Create a sample DOCX file for testing."""
    file_path = temp_dir / "test.docx"
    sample_docx_content.save(str(file_path))
    return file_path

@pytest.mark.asyncio
async def test_docx_document_creation(document_processor: DocumentProcessor, test_user: Dict, sample_docx_file: Path):
    """Test creating a new DOCX document."""
    doc = await document_processor.create_document(
        user_id=test_user["id"],
        filename="test.docx",
        file_path=str(sample_docx_file)
    )
    
    assert isinstance(doc, Document)
    assert doc.user_id == test_user["id"]
    assert doc.filename == "test.docx"
    assert doc.status == DocumentStatus.PENDING

@pytest.mark.asyncio
async def test_docx_processing(document_processor: DocumentProcessor, test_user: Dict, sample_docx_file: Path):
    """Test processing a DOCX document."""
    # Create document
    doc = await document_processor.create_document(
        user_id=test_user["id"],
        filename="test.docx",
        file_path=str(sample_docx_file)
    )
    
    # Process document
    processed_doc = await document_processor.process_document(doc)
    
    assert processed_doc.status == DocumentStatus.COMPLETED
    assert processed_doc.metadata.get("user_id") == test_user["id"]
    assert processed_doc.metadata.get("document_name") == "test.docx"

@pytest.mark.asyncio
async def test_docx_chunk_generation(document_processor: DocumentProcessor, test_user: Dict, sample_docx_file: Path):
    """Test generating chunks from a DOCX document."""
    # Create document
    doc = await document_processor.create_document(
        user_id=test_user["id"],
        filename="test.docx",
        file_path=str(sample_docx_file)
    )
    
    # Generate chunks
    chunks = await document_processor.generate_chunks(doc)
    
    assert len(chunks) == 1  # We combine all paragraphs into one chunk
    chunk = chunks[0]
    assert chunk.document_id == doc.id
    assert chunk.user_id == doc.user_id
    assert isinstance(chunk.content, str)
    assert len(chunk.content) > 0
    assert "Test Content" in chunk.content
    assert "Second paragraph" in chunk.content
    assert chunk.metadata["source"] == "docx"

@pytest.mark.asyncio
async def test_docx_error_handling(document_processor: DocumentProcessor, test_user: Dict, temp_dir: Path):
    """Test error handling for DOCX processing."""
    # Create document with non-existent file
    doc = await document_processor.create_document(
        user_id=test_user["id"],
        filename="nonexistent.docx",
        file_path=str(temp_dir / "nonexistent.docx")
    )
    
    # Process document should handle the error
    processed_doc = await document_processor.process_document(doc)
    
    assert processed_doc.status == DocumentStatus.ERROR
    assert processed_doc.error_message is not None

@pytest.mark.asyncio
async def test_docx_size_limit(document_processor: DocumentProcessor, test_user: Dict, temp_dir: Path):
    """Test file size limit for DOCX processing."""
    # Create a large DOCX file that exceeds the size limit
    large_file = temp_dir / "large.docx"
    doc = DocxDocument()
    # Add enough content to exceed the size limit
    for _ in range(1000000):  # This should create a file larger than 100MB
        doc.add_paragraph("Large content " * 100)
    doc.save(str(large_file))
    
    # Create document
    doc = await document_processor.create_document(
        user_id=test_user["id"],
        filename="large.docx",
        file_path=str(large_file)
    )
    
    # Process document should fail with size error
    processed_doc = await document_processor.process_document(doc)
    
    assert processed_doc.status == DocumentStatus.ERROR
    assert "size" in processed_doc.error_message.lower() 