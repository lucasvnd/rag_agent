import os
import pytest
from docx import Document
from docx.document import Document as DocumentType
from src.models.template import Template
from src.services.template_processor import TemplateProcessor

@pytest.fixture
def template_dir():
    return os.path.join(os.path.dirname(os.path.dirname(__file__)), 'templates')

@pytest.fixture
def template_files(template_dir):
    return [f for f in os.listdir(template_dir) if f.endswith('.docx')]

def test_template_files_exist(template_dir):
    """Test that template directory exists and contains .docx files"""
    assert os.path.exists(template_dir)
    files = os.listdir(template_dir)
    docx_files = [f for f in files if f.endswith('.docx')]
    assert len(docx_files) > 0, "No .docx files found in templates directory"

def test_template_file_structure(template_dir, template_files):
    """Test that each template file can be opened and has valid structure"""
    for file_name in template_files:
        file_path = os.path.join(template_dir, file_name)
        assert os.path.exists(file_path), f"Template file {file_name} does not exist"
        
        # Try to open the file with python-docx
        try:
            doc = Document(file_path)
            # Basic structure checks
            assert len(doc.paragraphs) > 0, f"Template {file_name} has no paragraphs"
        except Exception as e:
            pytest.fail(f"Failed to open template {file_name}: {str(e)}")

def test_template_variables(template_dir, template_files):
    """Test that template variables can be extracted and are valid"""
    processor = TemplateProcessor()
    
    for file_name in template_files:
        file_path = os.path.join(template_dir, file_name)
        template = Template(
            name=file_name,
            path=file_path,
            description="Test template"
        )
        
        # Extract variables
        variables = processor.extract_variables(template)
        
        # Verify variables
        assert isinstance(variables, list), f"Variables for {file_name} should be a list"
        # Note: Some templates might not have variables, so we don't assert length > 0
        
        # Check variable format if any exist
        for var in variables:
            assert isinstance(var, str), f"Variable {var} in {file_name} is not a string"
            assert var.strip() == var, f"Variable {var} in {file_name} has leading/trailing whitespace"
            assert len(var) > 0, f"Empty variable found in {file_name}"

def test_template_processing(template_dir, template_files):
    """Test that templates can be processed with sample data"""
    processor = TemplateProcessor()
    
    for file_name in template_files:
        file_path = os.path.join(template_dir, file_name)
        template = Template(
            name=file_name,
            path=file_path,
            description="Test template"
        )
        
        # Get variables
        variables = processor.extract_variables(template)
        
        # Create sample data
        sample_data = {var: f"Sample_{var}" for var in variables}
        
        # Try processing
        try:
            result = processor.process_template(template, sample_data)
            assert result is not None, f"Processing {file_name} returned None"
            assert isinstance(result, DocumentType), f"Processing {file_name} did not return a Document"
        except Exception as e:
            pytest.fail(f"Failed to process template {file_name}: {str(e)}")

def test_template_metadata(template_dir, template_files):
    """Test template metadata extraction and validation"""
    for file_name in template_files:
        file_path = os.path.join(template_dir, file_name)
        doc = Document(file_path)
        
        # Check core properties
        core_props = doc.core_properties
        assert core_props is not None, f"No core properties found in {file_name}"
        
        # Basic metadata checks - these might be empty but should exist
        assert hasattr(core_props, 'author'), f"No author property in {file_name}"
        assert hasattr(core_props, 'title'), f"No title property in {file_name}" 