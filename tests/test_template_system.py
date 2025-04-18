import pytest
from pathlib import Path
import uuid
from datetime import datetime
from typing import Dict, Any
from docx import Document

from models.templates import Template, TemplateCreate, TemplateUpdate
from services.template_processor import TemplateProcessor

# Test data
TEST_TEMPLATE_PATH = Path("templates/modelo_a.docx")
TEST_OUTPUT_PATH = Path("test_output")

@pytest.fixture
def template_processor():
    return TemplateProcessor()

@pytest.fixture
def sample_template():
    return Template(
        id=uuid.uuid4(),
        user_id=uuid.uuid4(),
        name="test_template",
        description="Test template",
        file_path=str(TEST_TEMPLATE_PATH),
        metadata={},
        version=1,
        is_active=True,
        created_at=datetime.now(),
        updated_at=datetime.now()
    )

@pytest.mark.asyncio
async def test_template_loading(template_processor: TemplateProcessor, sample_template: Template):
    """Test loading a template file"""
    doc = await template_processor.load_template(sample_template)
    assert isinstance(doc, Document)

@pytest.mark.asyncio
async def test_variable_extraction(template_processor: TemplateProcessor, sample_template: Template):
    """Test extracting variables from template"""
    variables = await template_processor.get_template_variables(sample_template)
    assert isinstance(variables, dict)
    # We should find at least one variable in the template
    assert len(variables) > 0
    # Variables should be in {{variable}} format
    for var in variables.keys():
        assert "{{" not in var and "}}" not in var

@pytest.mark.asyncio
async def test_template_validation(template_processor: TemplateProcessor, sample_template: Template):
    """Test template data validation"""
    # First get the required variables
    variables = await template_processor.get_template_variables(sample_template)
    
    # Test with missing variables
    test_data: Dict[str, Any] = {}
    with pytest.raises(Exception) as exc_info:
        await template_processor.validate_template_data(sample_template, test_data)
    assert "Missing required variables" in str(exc_info.value)
    
    # Test with all variables provided
    test_data = {var: f"test_{var}" for var in variables.keys()}
    result = await template_processor.validate_template_data(sample_template, test_data)
    assert result is True

@pytest.mark.asyncio
async def test_template_processing(template_processor: TemplateProcessor, sample_template: Template):
    """Test processing a template with variables"""
    # Get variables and create test data
    variables = await template_processor.get_template_variables(sample_template)
    test_data = {var: f"test_{var}" for var in variables.keys()}
    
    # Process template
    doc = await template_processor.process_template(sample_template, test_data)
    assert isinstance(doc, Document)
    
    # Check if variables were replaced
    for paragraph in doc.paragraphs:
        assert "{{" not in paragraph.text and "}}" not in paragraph.text
    
    # Check tables if they exist
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    assert "{{" not in paragraph.text and "}}" not in paragraph.text

@pytest.mark.asyncio
async def test_save_processed_template(template_processor: TemplateProcessor, sample_template: Template):
    """Test saving a processed template"""
    # Ensure output directory exists
    TEST_OUTPUT_PATH.mkdir(exist_ok=True)
    
    # Process template
    variables = await template_processor.get_template_variables(sample_template)
    test_data = {var: f"test_{var}" for var in variables.keys()}
    doc = await template_processor.process_template(sample_template, test_data)
    
    # Save processed template
    output_path = str(TEST_OUTPUT_PATH / "processed_template.docx")
    saved_path = await template_processor.save_processed_template(doc, output_path)
    
    assert Path(saved_path).exists()
    assert Path(saved_path).is_file()

@pytest.mark.asyncio
async def test_template_metadata_update(template_processor: TemplateProcessor, sample_template: Template):
    """Test updating template metadata with variables"""
    variables = await template_processor.get_template_variables(sample_template)
    metadata = await template_processor.update_template_metadata(sample_template, variables)
    
    assert "variables" in metadata
    assert isinstance(metadata["variables"], dict)
    assert len(metadata["variables"]) == len(variables)

def test_cleanup():
    """Clean up test output files"""
    if TEST_OUTPUT_PATH.exists():
        for file in TEST_OUTPUT_PATH.glob("*.docx"):
            file.unlink()
        TEST_OUTPUT_PATH.rmdir() 